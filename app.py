import re
from io import BytesIO
import streamlit as st
from docx import Document
from collections import defaultdict, Counter

st.set_page_config(page_title="DOCX: Interactive Chapter Detection & Citation Processor", page_icon="📚")
st.title("Interactive Font-Based Chapter Detection & Citation Processing")

st.markdown("""
### Interactive Workflow:
1. **📊 Analyze**: Examine all font sizes in your document
2. **🎯 Choose**: Select which font size to use for chapter detection  
3. **✅ Select**: Pick specific paragraphs as chapter headers
4. **⚙️ Process**: Handle citations in each chapter
5. **📥 Download**: Get processed document
""")

uploaded = st.file_uploader("Upload your DOCX file", type=["docx"])

def analyze_document_fonts(doc):
    """Comprehensive font analysis of the document"""
    font_analysis = {
        'font_sizes': Counter(),
        'font_examples': defaultdict(list),
        'total_paragraphs': len(doc.paragraphs),
        'empty_paragraphs': 0
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            font_analysis['empty_paragraphs'] += 1
            continue
        
        # Analyze each run in paragraph
        paragraph_fonts = []
        for run in para.runs:
            if run.font.size:
                font_size = run.font.size.pt
                paragraph_fonts.append(font_size)
                font_analysis['font_sizes'][font_size] += 1
        
        # Store example text for each font size found in this paragraph
        if paragraph_fonts:
            # Use the largest font in the paragraph as the representative size
            max_font = max(paragraph_fonts)
            if len(font_analysis['font_examples'][max_font]) < 10:  # Limit examples
                font_analysis['font_examples'][max_font].append({
                    'paragraph_index': i,
                    'text': text[:100] + ('...' if len(text) > 100 else ''),
                    'full_text': text,
                    'font_size': max_font
                })
    
    return font_analysis

def display_font_analysis(analysis):
    """Display font analysis results"""
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Total Paragraphs", analysis['total_paragraphs'])
    with col2:
        st.metric("Non-empty Paragraphs", analysis['total_paragraphs'] - analysis['empty_paragraphs'])
    with col3:
        st.metric("Unique Font Sizes", len(analysis['font_sizes']))
    
    # Display font size distribution
    st.subheader("📊 Font Size Distribution")
    
    if analysis['font_sizes']:
        # Create chart data
        font_data = dict(analysis['font_sizes'].most_common())
        st.bar_chart(font_data)
        
        # Display detailed breakdown
        st.subheader("📝 Font Size Examples")
        
        # Sort font sizes in descending order
        sorted_fonts = sorted(analysis['font_sizes'].items(), key=lambda x: x[0], reverse=True)
        
        for font_size, count in sorted_fonts:
            with st.expander(f"Font Size {font_size}pt ({count} occurrences)"):
                examples = analysis['font_examples'][font_size]
                for example in examples[:5]:  # Show max 5 examples
                    st.write(f"• **Para {example['paragraph_index']}**: {example['text']}")
    
    return analysis['font_sizes']

def get_chapter_candidates(doc, selected_font_size):
    """Get all paragraphs with the selected font size as chapter candidates"""
    candidates = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Check if paragraph has the selected font size
        has_target_font = False
        for run in para.runs:
            if run.font.size and run.font.size.pt == selected_font_size:
                has_target_font = True
                break
        
        if has_target_font:
            candidates.append({
                'index': i,
                'text': text,
                'preview': text[:80] + ('...' if len(text) > 80 else '')
            })
    
    return candidates

def create_chapter_boundaries(selected_chapters, total_paragraphs):
    """Create chapter boundaries from selected chapter headers"""
    if not selected_chapters:
        return [(0, total_paragraphs - 1, "Full_Document")]
    
    # Sort selected chapters by index
    selected_chapters.sort(key=lambda x: x['index'])
    
    boundaries = []
    for i, chapter in enumerate(selected_chapters):
        start = chapter['index']
        end = selected_chapters[i + 1]['index'] - 1 if i + 1 < len(selected_chapters) else total_paragraphs - 1
        
        # Clean title for filename
        title = re.sub(r'[^\w\s-]', '', chapter['text'])
        title = re.sub(r'\s+', '_', title)[:50]
        if not title:
            title = f"Chapter_{i+1}"
        
        boundaries.append((start, end, title))
    
    return boundaries

# Citation processing functions (keeping the same as before)
HEADING_RE = re.compile(r"^\s*(notes?|references?|endnotes?|sources?|bibliography|citations?)\s*:?\s*$", re.I)

def para_iter(doc):
    """Iterate through all paragraphs including tables"""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield p

def is_notes_heading(p):
    text = (p.text or "").strip()
    return bool(HEADING_RE.match(text)) if text else False

def find_notes_sections(paragraphs):
    sections = []
    for i, p in enumerate(paragraphs):
        if is_notes_heading(p):
            start = i + 1
            end = find_section_end(paragraphs, start)
            sections.append((i, start, end))
    return sections

def find_section_end(paragraphs, start):
    consecutive_blanks = 0
    for i in range(start, len(paragraphs)):
        p = paragraphs[i]
        text = (p.text or "").strip()
        if is_notes_heading(p):
            return i
        if not text:
            consecutive_blanks += 1
            if consecutive_blanks >= 2:
                return i
        else:
            consecutive_blanks = 0
    return len(paragraphs)

def parse_references_advanced(paragraphs, start, end):
    refs = {}
    current_num = None
    current_text = ""
    
    for i in range(start, end):
        if i >= len(paragraphs):
            break
        p = paragraphs[i]
        text = (p.text or "").strip()
        if not text:
            continue
        
        patterns = [
            r"^(\d+)\.\s*(.+)$",
            r"^(\d+)\)\s*(.+)$", 
            r"^(\d+)\]\s*(.+)$",
            r"^(\d+)\s+(.+)$",
            r"^(\d+)[\-–—:]\s*(.+)$",
        ]
        
        found_match = False
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                if current_num is not None and current_text.strip():
                    refs[current_num] = current_text.strip()
                current_num = int(match.group(1))
                current_text = match.group(2)
                found_match = True
                break
        
        if not found_match and current_num is not None:
            current_text += " " + text
    
    if current_num is not None and current_text.strip():
        refs[current_num] = current_text.strip()
    
    return refs

def expand_number_range(num_str):
    nums = []
    parts = re.split(r'[,\s]+', num_str.replace('–', '-').replace('—', '-'))
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if '-' in part:
            try:
                start, end = part.split('-', 1)
                start, end = int(start), int(end)
                if 1 <= start <= end <= 999 and (end - start) < 50:
                    nums.extend(range(start, end + 1))
            except:
                pass
        elif part.isdigit():
            nums.append(int(part))
    return sorted(set(nums))

def format_reference(num, text, style):
    if style.startswith("—"):
        return f" — {num}. {text}"
    elif style.startswith("["):
        return f" [{num}. {text}]"
    else:
        return f" ({text})"

def replace_citations_in_paragraph(paragraph, all_refs, style, allow_paren=False):
    if not all_refs:
        return 0
    
    changes = 0
    max_ref_num = max(all_refs.keys())
    
    # Handle superscript runs
    for run in paragraph.runs:
        if getattr(run.font, 'superscript', None):
            nums = [int(x) for x in re.findall(r'\d+', run.text)]
            if (nums and all(1 <= n <= max_ref_num and n in all_refs for n in nums) 
                and not any(n >= 1000 for n in nums)):
                replacements = [format_reference(n, all_refs[n], style).strip() for n in nums]
                run.font.superscript = None
                run.text = "; ".join(replacements)
                changes += 1
    
    # Handle text-based citations
    original_text = paragraph.text
    modified_text = original_text
    
    def replace_square(match):
        nums = expand_number_range(match.group(1))
        if (nums and all(1 <= n <= max_ref_num and n in all_refs for n in nums) 
            and not any(n >= 1000 for n in nums)):
            replacements = [format_reference(n, all_refs[n], style) for n in nums]
            return "; ".join(replacements)
        return match.group(0)
    
    modified_text = re.sub(r'\[([0-9,\-–—\s]+)\]', replace_square, modified_text)
    
    if allow_paren:
        modified_text = re.sub(r'\(([0-9,\-–—\s]+)\)', replace_square, modified_text)
    
    if modified_text != original_text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = modified_text
        else:
            paragraph.add_run(modified_text)
        changes += 1
    
    return changes

def create_chapter_document(original_doc, start_para, end_para):
    new_doc = Document()
    for i in range(start_para, min(end_para + 1, len(original_doc.paragraphs))):
        old_para = original_doc.paragraphs[i]
        new_para = new_doc.add_paragraph()
        try:
            new_para.style = old_para.style
        except:
            pass
        for run in old_para.runs:
            new_run = new_para.add_run(run.text)
            try:
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
                new_run.font.underline = run.font.underline
                if hasattr(run.font, 'superscript'):
                    new_run.font.superscript = run.font.superscript
            except:
                pass
    return new_doc

def process_single_chapter(doc, style, allow_paren=False, delete_notes_sections=False):
    paragraphs = list(para_iter(doc))
    notes_sections = find_notes_sections(paragraphs)
    
    if not notes_sections:
        return 0, 0
    
    all_refs = {}
    for section_head, section_start, section_end in notes_sections:
        refs = parse_references_advanced(paragraphs, section_start, section_end)
        all_refs.update(refs)
    
    if not all_refs:
        return 0, 0
    
    total_replacements = 0
    notes_ranges = set()
    for section_head, section_start, section_end in notes_sections:
        for i in range(section_head, section_end):
            notes_ranges.add(i)
    
    for i, paragraph in enumerate(paragraphs):
        if i not in notes_ranges and paragraph.text and paragraph.text.strip():
            replacements = replace_citations_in_paragraph(paragraph, all_refs, style, allow_paren)
            total_replacements += replacements
    
    if delete_notes_sections:
        for section_head, section_start, section_end in reversed(notes_sections):
            for idx in range(section_end - 1, section_head - 1, -1):
                if idx < len(paragraphs):
                    p = paragraphs[idx]
                    if p._element.getparent() is not None:
                        p._element.getparent().remove(p._element)
    
    return len(all_refs), total_replacements

def rejoin_chapters(chapter_docs):
    final_doc = Document()
    for i, chapter_doc in enumerate(chapter_docs):
        if i > 0:
            final_doc.add_page_break()
        for para in chapter_doc.paragraphs:
            new_para = final_doc.add_paragraph()
            try:
                new_para.style = para.style
            except:
                pass
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                try:
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
                    new_run.font.underline = run.font.underline
                except:
                    pass
    return final_doc

# Main application
if uploaded:
    try:
        doc = Document(uploaded)
        
        # Step 1: Analyze document fonts
        st.header("📊 Step 1: Font Analysis")
        with st.spinner("Analyzing document fonts..."):
            analysis = analyze_document_fonts(doc)
        
        font_sizes = display_font_analysis(analysis)
        
        if not font_sizes:
            st.error("No font sizes detected in the document!")
            st.stop()
        
        # Step 2: Let user choose font size for chapter detection
        st.header("🎯 Step 2: Choose Chapter Detection Font Size")
        
        available_sizes = sorted(font_sizes.keys(), reverse=True)
        selected_font_size = st.selectbox(
            "Select font size for chapter detection:",
            available_sizes,
            format_func=lambda x: f"{x}pt ({font_sizes[x]} occurrences)"
        )
        
        # Step 3: Show candidates and let user select chapters
        st.header("✅ Step 3: Select Chapter Headers")
        
        candidates = get_chapter_candidates(doc, selected_font_size)
        
        if not candidates:
            st.warning(f"No paragraphs found with font size {selected_font_size}pt")
            st.stop()
        
        st.write(f"Found **{len(candidates)}** paragraphs with font size {selected_font_size}pt:")
        
        # Let user select which candidates should be chapter headers
        selected_indices = []
        
        for i, candidate in enumerate(candidates):
            is_selected = st.checkbox(
                f"**Para {candidate['index']}**: {candidate['preview']}",
                key=f"chapter_{i}",
                help=f"Full text: {candidate['text']}"
            )
            if is_selected:
                selected_indices.append(i)
        
        if not selected_indices:
            st.warning("Please select at least one paragraph as a chapter header.")
            st.stop()
        
        selected_chapters = [candidates[i] for i in selected_indices]
        
        # Step 4: Show chapter boundaries
        st.header("📋 Step 4: Chapter Boundaries Preview")
        
        boundaries = create_chapter_boundaries(selected_chapters, len(doc.paragraphs))
        
        for i, (start, end, title) in enumerate(boundaries):
            st.write(f"**{i+1:02d}. {title}** - Paragraphs {start} to {end} ({end-start+1} paragraphs)")
        
        # Step 5: Processing options and execution
        st.header("⚙️ Step 5: Processing Options")
        
        col1, col2 = st.columns(2)
        with col1:
            fmt = st.selectbox("Citation format:", ["[1. Reference text]", "— 1. Reference text", "(Reference text)"])
        with col2:
            delete_notes = st.checkbox("Delete Notes sections after processing")
            allow_paren = st.checkbox("Also convert (n) patterns (risky near years)")
        
        if st.button("🚀 Process All Chapters", type="primary"):
            chapter_docs = []
            total_refs = 0
            total_replacements = 0
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Process each chapter
            for i, (start, end, title) in enumerate(boundaries):
                status_text.text(f"Processing Chapter {i+1}/{len(boundaries)}: {title}")
                
                # Create and process chapter
                chapter_doc = create_chapter_document(doc, start, end)
                refs_found, citations_replaced = process_single_chapter(
                    chapter_doc, fmt, allow_paren, delete_notes
                )
                
                chapter_docs.append(chapter_doc)
                total_refs += refs_found
                total_replacements += citations_replaced
                
                st.write(f"✅ **{i+1:02d}. {title}**: {refs_found} references, {citations_replaced} citations replaced")
                progress_bar.progress((i + 1) / len(boundaries))
            
            # Rejoin chapters
            status_text.text("🔗 Rejoining all processed chapters...")
            final_doc = rejoin_chapters(chapter_docs)
            
            # Save and offer download
            bio = BytesIO()
            final_doc.save(bio)
            bio.seek(0)
            
            status_text.text("✅ Processing complete!")
            st.success(f"🎉 **Final Results**: {total_refs} total references, {total_replacements} total citations replaced across {len(boundaries)} chapters!")
            
            st.download_button(
                "📥 Download Processed Document",
                data=bio.getvalue(),
                file_name="book_processed_interactive.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Optional: Download individual chapters
            with st.expander("📁 Download Individual Chapters"):
                for i, (chapter_doc, (start, end, title)) in enumerate(zip(chapter_docs, boundaries)):
                    chapter_bio = BytesIO()
                    chapter_doc.save(chapter_bio)
                    chapter_bio.seek(0)
                    
                    st.download_button(
                        f"Chapter {i+1:02d}: {title}",
                        data=chapter_bio.getvalue(),
                        file_name=f"chapter_{i+1:02d}_{title}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"chapter_{i}"
                    )

    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        st.exception(e)

else:
    st.info("📤 Please upload a DOCX file to begin interactive analysis")
    st.markdown("""
    ### How This Works:
    1. **📊 Font Analysis**: Shows all font sizes in your document with examples
    2. **🎯 Font Selection**: You choose which font size indicates chapter headers  
    3. **✅ Chapter Selection**: Preview and select specific paragraphs as chapters
    4. **⚙️ Processing**: Each chapter processed individually for better accuracy
    5. **📥 Results**: Download processed document plus individual chapters
    
    ### Benefits:
    - ✅ **Full Control**: You decide what constitutes a chapter
    - ✅ **Visual Preview**: See examples before making decisions
    - ✅ **Flexible**: Works with any font size or document structure
    - ✅ **Accurate**: Chapter-wise processing improves citation detection
    """)
