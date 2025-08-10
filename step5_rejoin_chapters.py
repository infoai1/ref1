import streamlit as st
from docx import Document
from io import BytesIO

def rejoin_chapters_with_formatting(chapter_docs):
    """Rejoin chapters preserving exact formatting"""
    final_doc = Document()
    
    for i, chapter_doc in enumerate(chapter_docs):
        # Add page break between chapters (except first)
        if i > 0:
            final_doc.add_page_break()
        
        # Copy all paragraphs with exact formatting
        for para in chapter_doc.paragraphs:
            new_para = final_doc.add_paragraph()
            
            # Copy paragraph style
            try:
                new_para.style = para.style
            except:
                pass
            
            # Copy all runs with exact formatting
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                
                # Preserve ALL formatting
                try:
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
                    new_run.font.underline = run.font.underline
                    if hasattr(run.font, 'color') and run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
                    if hasattr(run.font, 'superscript'):
                        new_run.font.superscript = run.font.superscript
                    if hasattr(run.font, 'subscript'):
                        new_run.font.subscript = run.font.subscript
                except:
                    pass
    
    return final_doc

# Only run UI code if this file is run directly
if __name__ == "__main__":
    st.set_page_config(page_title="Step 5: Rejoin Chapters", page_icon="ğŸ”—")
    st.title("ğŸ”— Step 5: Rejoin Processed Chapters")
    
    processed_data = st.session_state.get('processed_chapters')
    
    if processed_data:
        chapter_docs = processed_data['chapter_docs']
        boundaries = processed_data['boundaries']
        stats = processed_data['stats']
        
        st.subheader("ğŸ“‹ Processing Summary")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Chapters Processed", len(chapter_docs))
        with col2:
            st.metric("References Found", stats['refs'])
        with col3:
            st.metric("Citations Replaced", stats['replacements'])
        
        st.subheader("ğŸ“– Chapter List")
        for i, (start, end, title) in enumerate(boundaries):
            st.write(f"**{i+1:02d}. {title}** (Paragraphs {start}-{end})")
        
        if st.button("ğŸ”— Rejoin All Chapters", type="primary"):
            with st.spinner("Rejoining chapters with preserved formatting..."):
                final_doc = rejoin_chapters_with_formatting(chapter_docs)
            
            # Save to BytesIO
            bio = BytesIO()
            final_doc.save(bio)
            bio.seek(0)
            
            st.success("âœ… Chapters successfully rejoined with preserved formatting!")
            
            # Download button
            st.download_button(
                "ğŸ“¥ Download Final Processed Document",
                data=bio.getvalue(),
                file_name="book_processed_final.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Optional: Download individual chapters
            with st.expander("ğŸ“ Download Individual Chapters"):
                for i, (chapter_doc, (start, end, title)) in enumerate(zip(chapter_docs, boundaries)):
                    chapter_bio = BytesIO()
                    chapter_doc.save(chapter_bio)
                    chapter_bio.seek(0)
                    
                    st.download_button(
                        f"Chapter {i+1:02d}: {title}",
                        data=chapter_bio.getvalue(),
                        file_name=f"chapter_{i+1:02d}_{title}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"ch_{i}"
                    )
    
    else:
        st.warning("âš ï¸ Please complete Step 4 first (Process Citations)")
        st.info("Run: `streamlit run step4_citation_processing.py`")
    
    st.markdown("---")
    st.info("ğŸ‰ **Workflow Complete!** Your document is processed with exact formatting preservation.")
