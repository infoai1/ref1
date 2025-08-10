import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from collections import Counter, defaultdict

def detect_all_font_sizes(doc):
    """Enhanced font size detection - finds ALL font sizes including hidden ones"""
    font_sizes = Counter()
    font_examples = defaultdict(list)
    
    # Method 1: Check document styles first (this often contains 26pt fonts)
    style_fonts = {}
    for style in doc.styles:
        try:
            if hasattr(style, 'font') and style.font:
                if style.font.size:
                    size_pt = style.font.size.pt
                    style_fonts[style.name] = size_pt
                    font_sizes[size_pt] += 1
        except:
            pass
    
    # Method 2: Check document default font
    try:
        doc_defaults = doc.part.element.xpath('.//w:docDefaults')[0]
        run_defaults = doc_defaults.xpath('.//w:rPrDefault')
        sz_elem = run_defaults.xpath('.//w:sz')
        if sz_elem:
            default_size = float(sz_elem.get(qn('w:val'))) / 2
            font_sizes[default_size] += 1
    except:
        pass
    
    # Method 3: Deep paragraph analysis with multiple detection methods
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Get style font size for this paragraph
        style_font_size = None
        try:
            style_name = para.style.name
            style_font_size = style_fonts.get(style_name)
        except:
            pass
        
        # Check each run with multiple methods
        para_font_sizes = []
        
        for run in para.runs:
            font_size = style_font_size  # Start with style font
            
            # Method 3a: Direct font.size property
            try:
                if run.font.size:
                    font_size = run.font.size.pt
            except:
                pass
            
            # Method 3b: XML-level font size detection
            try:
                if run.element.rPr is not None:
                    # Check w:sz element (font size in half-points)
                    sz = run.element.rPr.find(qn('w:sz'))
                    if sz is not None:
                        font_size = float(sz.get(qn('w:val'))) / 2
                    
                    # Check w:szCs element (complex script font size)
                    szCs = run.element.rPr.find(qn('w:szCs'))
                    if szCs is not None:
                        font_size = max(font_size or 0, float(szCs.get(qn('w:val'))) / 2)
            except:
                pass
            
            # Method 3c: Check for theme fonts and complex scripts
            try:
                if run.element.rPr is not None:
                    # Sometimes fonts are defined in theme
                    rFonts = run.element.rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        # Check various font attributes
                        ascii_theme = rFonts.get(qn('w:asciiTheme'))
                        if ascii_theme and not font_size:
                            font_size = 12  # Default theme size, but check further
            except:
                pass
            
            if font_size:
                para_font_sizes.append(font_size)
        
        # Method 4: Check paragraph-level formatting
        try:
            if para.paragraph_format and hasattr(para.paragraph_format, 'font'):
                if para.paragraph_format.font and para.paragraph_format.font.size:
                    para_font_sizes.append(para.paragraph_format.font.size.pt)
        except:
            pass
        
        # Use the largest font size found in this paragraph
        if para_font_sizes:
            max_font = max(para_font_sizes)
            font_sizes[max_font] += 1
            
            # Store examples, especially for larger fonts
            if len(font_examples[max_font]) < 10 or max_font >= 20:
                font_examples[max_font].append({
                    'para_index': i,
                    'text': text[:80] + ('...' if len(text) > 80 else ''),
                    'full_text': text,
                    'method': 'paragraph_analysis'
                })
    
    # Method 5: XML document-wide search for font sizes
    try:
        # Search entire document XML for sz elements
        xml_content = doc.part.element.xml
        import re
        
        # Find all w:sz values in the document
        sz_pattern = r'w:sz w:val="(\d+)"'
        sz_matches = re.findall(sz_pattern, xml_content)
        
        for sz_val in sz_matches:
            size_pt = float(sz_val) / 2  # Convert half-points to points
            if size_pt not in font_sizes:
                font_sizes[size_pt] += 1
                font_examples[size_pt].append({
                    'para_index': -1,
                    'text': f'Found in XML: {size_pt}pt',
                    'full_text': f'Detected via XML parsing: {size_pt}pt font size',
                    'method': 'xml_search'
                })
    except:
        pass
    
    # Method 6: Check table cells (fonts might be hidden in tables)
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            for run in para.runs:
                                font_size = None
                                try:
                                    if run.font.size:
                                        font_size = run.font.size.pt
                                except:
                                    pass
                                
                                try:
                                    if run.element.rPr is not None:
                                        sz = run.element.rPr.find(qn('w:sz'))
                                        if sz is not None:
                                            font_size = float(sz.get(qn('w:val'))) / 2
                                except:
                                    pass
                                
                                if font_size:
                                    font_sizes[font_size] += 1
    except:
        pass
    
    return font_sizes, font_examples

# Only run UI code if this file is run directly
if __name__ == "__main__":
    st.set_page_config(page_title="Step 1: Enhanced Font Analysis", page_icon="🔍")
    st.title("🔍 Step 1: Enhanced Font Size Analysis")
    
    uploaded = st.file_uploader("Upload your DOCX file", type=["docx"])
    
    if uploaded:
        doc = Document(uploaded)
        
        st.subheader("📊 Analyzing Font Sizes (Enhanced Detection)...")
        font_sizes, font_examples = detect_all_font_sizes(doc)
        
        if font_sizes:
            st.success(f"Found {len(font_sizes)} different font sizes using enhanced detection!")
            
            # Display results
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Total Paragraphs", len(doc.paragraphs))
            with col2:
                st.metric("Font Sizes Found", len(font_sizes))
            
            # Highlight if font size 26 was found
            if 26.0 in font_sizes:
                st.success("🎉 **Font size 26pt detected!** ✅")
            else:
                st.warning("⚠️ Font size 26pt not found in this document")
            
            # Chart
            st.subheader("📈 Font Size Distribution")
            chart_data = dict(font_sizes.most_common())
            st.bar_chart(chart_data)
            
            # Examples with detection method
            st.subheader("📝 Font Size Examples")
            for font_size, count in sorted(font_sizes.items(), reverse=True):
                # Highlight larger fonts
                color = "🔴" if font_size >= 24 else "🟡" if font_size >= 18 else "🟢"
                
                with st.expander(f"{color} Font Size {font_size}pt ({count} occurrences)"):
                    examples = font_examples[font_size]
                    for ex in examples[:5]:  # Show max 5 examples
                        if ex['para_index'] >= 0:
                            st.write(f"• **Para {ex['para_index']}**: {ex['text']}")
                        else:
                            st.write(f"• **{ex['text']}** (Method: {ex.get('method', 'unknown')})")
            
            # Save results for next step
            st.session_state['font_analysis'] = {
                'font_sizes': dict(font_sizes),
                'font_examples': dict(font_examples),
                'doc_path': uploaded.name
            }
            
            st.success("✅ Enhanced analysis complete! Proceed to Step 2.")
            
            # Show summary of large fonts found
            large_fonts = [size for size in font_sizes.keys() if size >= 20]
            if large_fonts:
                st.info(f"📊 **Large fonts detected**: {sorted(large_fonts, reverse=True)}")
            
        else:
            st.error("No font sizes detected even with enhanced detection!")
            st.info("This might indicate an issue with the document structure.")
    else:
        st.info("📤 Upload a DOCX file to analyze font sizes")
        st.markdown("""
        ### 🔍 Enhanced Detection Methods:
        1. **Document Styles** - Checks style definitions
        2. **XML Parsing** - Direct XML font size extraction  
        3. **Run Analysis** - Individual text run fonts
        4. **Theme Fonts** - Document theme-based fonts
        5. **Table Fonts** - Fonts hidden in table cells
        6. **Complex Scripts** - Multi-language font support
        
        This should detect **ALL** font sizes including 26pt!
        """)
