import streamlit as st
from docx import Document
from io import BytesIO
import re

# Citation processing functions
HEADING_RE = re.compile(r"^\s*(notes?|references?|endnotes?|sources?|bibliography|citations?)\s*:?\s*$", re.I)

def para_iter(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield p

def find_notes_sections(paragraphs):
    sections = []
    for i, p in enumerate(paragraphs):
        if p.text and HEADING_RE.match(p.text.strip()):
            start = i + 1
            end = find_section_end(paragraphs, start)
            sections.append((i, start, end))
    return sections

def find_section_end(paragraphs, start):
    consecutive_blanks = 0
    for i in range(start, len(paragraphs)):
        text = paragraphs[i].text.strip() if paragraphs[i].text else ""
        if not text:
            consecutive_blanks += 1
            if consecutive_blanks >= 2:
                return i
        else:
            consecutive_blanks = 0
    return len(paragraphs)

def parse_references(paragraphs, start, end):
    refs = {}
    current_num = None
    current_text = ""
    
    for i in range(start, end):
        if i >= len(paragraphs):
            break
        text = paragraphs[i].text.strip() if paragraphs[i].text else ""
        if not text:
            continue
        
        # Try to match numbered reference
        patterns = [
            r'^(\d+)[\.\)\]\s]+(.+)$',
            r'^(\d+)[\-–—:]\s*(.+)$'
        ]
        
        found_match = False
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                if current_num and current_text:
                    refs[current_num] = current_text.strip()
                current_num = int(match.group(1))
                current_text = match.group(2)
                found_match = True
                break
        
        if not found_match and current_num:
            current_text += " " + text
    
    if current_num and current_text:
        refs[current_num] = current_text.strip()
    
    return refs

def create_chapter_document(original_doc, start, end):
    """Create chapter document preserving ALL formatting"""
    new_doc = Document()
    
    for i in range(start, min(end + 1, len(original_doc.paragraphs))):
        old_para = original_doc.paragraphs[i]
        new_para = new_doc.add_paragraph()
        
        # Copy paragraph style
        try:
            new_para.style = old_para.style
        except:
            pass
        
        # Copy all runs with exact formatting
        for run in old_para.runs:
            new_run = new_para.add_run(run.text)
            
            # Copy ALL font properties
            try:
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
                new_run.font.underline = run.font.underline
                if hasattr(run.font, 'color') and run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
                if hasattr(run.font, 'superscript'):
                    new_run.font.superscript = run.font.superscript
                if hasattr(run.font, 'subscript'):
                    new_run.font.subscript = run.font.subscript
            except:
                pass
    
    return new_doc

def process_chapter_citations(doc, fmt="[1. Reference text]", delete_notes=False):
    """Process citations in a chapter while preserving formatting"""
    paragraphs = list(para_iter(doc))
    notes_sections = find_notes_sections(paragraphs)
    
    if not notes_sections:
        return 0, 0
    
    # Parse references
    all_refs = {}
    for _, start, end in notes_sections:
        refs = parse_references(paragraphs, start, end)
        all_refs.update(refs)
    
    if not all_refs:
        return 0, 0
    
    # Replace citations
    replacements = 0
    notes_ranges = set()
    for head, start, end in notes_sections:
        for i in range(head, end):
            notes_ranges.add(i)
    
    for i, para in enumerate(paragraphs):
        if i not in notes_ranges and para.text:
            original = para.text
            modified = original
            
            # Replace [1], [1] etc.
            def replace_citation(match):
                try:
                    num = int(match.group(1))
                    if num in all_refs:
                        if fmt.startswith('['):
                            return f" [{num}. {all_refs[num]}]"
                        elif fmt.startswith('—'):
                            return f" — {num}. {all_refs[num]}"
                        else:
                            return f" ({all_refs[num]})"
                except:
                    pass
                return match.group(0)
            
            modified = re.sub(r'\[(\d+)\]', replace_citation, modified)
            
            if modified != original:
                # Update text preserving formatting
                if para.runs:
                    para.runs[0].text = modified
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run(modified)
                replacements += 1
    
    # Delete notes sections if requested
    if delete_notes:
        for head, start, end in reversed(notes_sections):
            for idx in range(end - 1, head - 1, -1):
                if idx < len(paragraphs):
                    p = paragraphs[idx]
                    if p._element.getparent() is not None:
                        p._element.getparent().remove(p._element)
    
    return len(all_refs), replacements

# Only run UI code if this file is run directly
if __name__ == "__main__":
    st.set_page_config(page_title="Step 4: Process Citations", page_icon="⚙️")
    st.title("⚙️ Step 4: Process Citations")
    
    uploaded = st.file_uploader("Upload your DOCX file", type=["docx"])
    
    if uploaded:
        final_chapters = st.session_state.get('final_chapters')
        
        if final_chapters:
            doc = Document(uploaded)
            boundaries = final_chapters['boundaries']
            
            st.subheader("⚙️ Processing Options")
            col1, col2 = st.columns(2)
            with col1:
                fmt = st.selectbox("Citation format:", ["[1. Reference text]", "— 1. Reference text", "(Reference text)"])
            with col2:
                delete_notes = st.checkbox("Delete Notes sections after processing")
            
            if st.button("🚀 Process All Chapters"):
                chapter_docs = []
                total_refs = 0
                total_replacements = 0
                
                progress = st.progress(0)
                status = st.empty()
                
                for i, (start, end, title) in enumerate(boundaries):
                    status.text(f"Processing {i+1}/{len(boundaries)}: {title}")
                    
                    # Create chapter with preserved formatting
                    chapter_doc = create_chapter_document(doc, start, end)
                    
                    # Process citations
                    refs, citations = process_chapter_citations(chapter_doc, fmt, delete_notes)
                    
                    chapter_docs.append(chapter_doc)
                    total_refs += refs
                    total_replacements += citations
                    
                    st.write(f"✅ **{title}**: {refs} references, {citations} citations replaced")
                    progress.progress((i + 1) / len(boundaries))
                
                # Save processed chapters
                st.session_state['processed_chapters'] = {
                    'chapter_docs': chapter_docs,
                    'boundaries': boundaries,
                    'stats': {'refs': total_refs, 'replacements': total_replacements}
                }
                
                status.text("✅ Processing complete!")
                st.success(f"🎉 Processed {len(boundaries)} chapters! {total_refs} references, {total_replacements} citations replaced")
                st.info("Proceed to Step 5 to rejoin chapters.")
                st.info("Run: `streamlit run step5_rejoin_chapters.py`")
        else:
            st.warning("⚠️ Please complete Step 3 first (Select Chapters)")
            st.info("Run: `streamlit run step3_chapter_selection.py`")
    else:
        st.info("📤 Upload the same DOCX file")
